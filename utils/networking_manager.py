from docx import Document
import os
import sys
import pyperclip

sys.path.append(sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath("utils")))))

from utils.contact_manager import tomide_adeoye_contact


def old_connection_unknown_person():
    return f"""Hello,

I trust that you are doing well. I am reaching out beacuse I wrote a scraper to search out top profile companies and recruiters looking to hire exceptional engineers on GitHub and my search led me to your profile. 

My name is Tomide; a software engineer  with Dukka, an African bookkeeping and payment processing company. Here I lead the development of frontend applications. My software have helped the company effectively track user signups and draw insights from different forms of userdata. Other softwares I have built include PicAndEdits, a software that leverages PicsArt APIs for image edits, a web app to monitor Unilever regulatory complaince and several web scrapers harvesting data from companies like PropertyPro and a lot more. 

I am also a software researcher and writer with a Blockchain Research organisations, where have written 4 technical papers on smart contracts and leveraged my data analytics skills to help the company track top performing research across their website.

I would speak with you, interview with you and identify how I could add value to your organisation. 

Please find below my portfolio and a copy of my CV.

I hope to hearing from you soon!

Best wishes,
"""

def person_frontend_engineer_application(contact_name, company, admiration):
    about_me =  f"""Dear {contact_name},
I hope you are having a great week.

My name is Tomide, and I am an engineer with a business background from NUTM (nutmng.org). I currently work with an African bookkeeping and payment processing company; my software tracks the company's customer signups, churn rate, ltv... I also function across the mobile and backend teams. Other interesting software I have built include (links in signature):

- PicAndEdits (React image transformation tool) 
- Bookkeeping mobile app (Flutter)
- Unicom: a compliance software for Unilever/JEE (React)
- Juice Spend: Payment startup landing pages (React)
- Python email leads generator and web scrapers. 

Outside of my engineering experience, I have done technical research with Blockchain organisations like the Smart Contract Research Forumn (Canada) and Web3Bridge. I have also started a startup and have been involved in the development of a few more while leveraging my business mentorships @ NUTM from minds like Soumitra Dutta (Dean, Oxford Business School) and Maurizio Caio (Partner, TLcom Capital) and more. My studies at NUTM, focusing on tech, strategy, and more, my engineering experience at Dukka, my research experience at SCRF, and my collaborations with diverse businesses (https://bit.ly/tomide-bio) make me a highly qualified candidate to function in multiple facets of a startup.

""" 
    about_company = f"""I admire how you {admiration}"""

    cta = f"""I would love to work with {company} and deploy the aggregate of my experiences toward skyrocketing your product development speed and profitability in 2023. 

Can we have a quick chat for me to show you what I can do and how I plan to add value to {company}?

Regards,
"""
    if admiration == "":
        about_company = ""

    return about_me + about_company + cta

def person_python_engineer_application(company):
    return f"""Dear Team,
I trust your year has been amazing.

I am Tomide, a python backend engineer and a data analyst at an African bookkeeping and payment processing company; Solutions I have built have helped my company track user signups and draw insights from different forms of userdata. Other softwares I have built include PicAndEdits, a software that leverages PicsArt APIs for AI powered image edits and several web scrapers targeting companies like PropertyPro and a lot more. 

My competencies are in python and you can find some of my open source projects here https://github.com/tomideadeoye/Python-scripts.

I would love to join your team and identify all ways I could help you build. I would appreciate a few minutes of your time to show you what I can do and how I plan to add value to {company}. 

Please find below my portfolio and a copy of my CV.

I can't wait to hear from you!

Best wishes,
"""



def connection_known_person(contact_name, role, company, company_mission, team):

    if team == True:
        contact_name = "Team"

    contact_name = contact_name.split("@")[0].split(".")[0].capitalize()
    return f"""Dear {contact_name},

I am Tomide, a software engineer working with Dukka, an African bookkeeping and payment processing company where I lead development of frontend applications. The application has helped the company effectively track user signups and draw insights from different forms of userdata. Other softwares I have built include PicAndEdits, a software that leverages PicsArt APIs for AI powered image edits and several web scrapers targeting companies like PropertyPro and a lot more. 

I am also a research fellow at the Smart Contract Research Forum (SCRF), where have written 4 technical research papers on smart contracts and leveraged my data analysis skills to help the company track top performing research across their website.

I saw a {role} opening at {company} and I found the role interesting. I especially relate to your vision of {company_mission}.

I would love to join your team and identify all ways I could help you build. I would appreciate a few minutes of your time to show you what I can do and how I plan to add value to {company}. 

Please find below my portfolio and a copy of my CV.
I look forward to hearing from you soon!

Best wishes,
"""

def connection_unknown_person():
    return f"""Hello Recruiter,
I trust that you are doing well. 

I am reaching out as regards the opening at your company for a junior software engineer. I am really excited about the role and I would love to join your team.

My name is Tomide Adeoye and I am a lead frontend engineer at an African payment processing company. While I lead frontend development, I also have expereinces across other technologies, particularly mobile engineering and web automation with python. My softwares help my company track user signups, visualize and draw insights from user data. Other softwares I have built range from image editing tools to a complaince web app and several web scrapers we use in harvesting data from competitor and a lot more. I am also a work in research, writing and documentation at a Blockchain Research organisation; I review research papers and leveraged my programming skills to help the company track and visualize top performing research content.

I would love to speak with you, identify how I could add value to your company and remain at the fore of your hiring list.

Please find below my portfolio and a copy of my CV, and let me know when it is a good time to speak. I hope to hear from you soon!

Best wishes,"""

def topic_of_interest():
    return F"""
Decentralized Finance!!
From a non-technical viewpoint, DeFi is interesting to discuss because of the enormous diversity of opinions about cryptocurrencies in general and DeFi. 
Many of the DeFi rebuttals focus on crypto's inability to bank the unbanked, I think this is very erroneous. Crypto definitely pulls down a lot of walls in terms of mobility of labor and financial inclusion. While it's true that some good folks here in Africa (myself included) don't have mobile devices or are unable to participate in DeFi, some can!!!
There are lots of Africans receiving remittances or paychecks via crypto because of the complexity of traditional means, and these people's spending contributes to their local economy instead of having emigration as the sole option for mobility of labor.
Another argument is often the efficiency of centralized systems; that's true in a lot of cases. But there are also innovations in DeFi that can't exist in traditional finance; the composability of DeFi apps comes to mind here - the sheer ease of building an application that interacts with another app. 
By simply knowing Solidity, I can write scripts that interact with Uniswap, Aave, and more simultaneously. This is, for example, the reason why flash loans exist. I can also write scripts that automate my savings and create investment portfolios designed to my exact specs. Nigerian investment apps like Cowrywise, and Piggevest, for example, merely present limited investable options. Smart Contracts on Ethereum lets me write the scripts and logic that power the investment strategy. 
Stablecoins and their backing (fiat, crypto, or algorithmic) are other topics worthy of discourse that I hope to continue during the interview. 
Read some of my DeFi business model deep-dive here: https://www.smartcontractresearch.org/t/research-summary-a-short-survey-on-business-models-of-decentralized-finance-defi-protocols/1947"""


def cool_project_worked_on():
    return F"""A project I am proud of is the Dexter Referral System.
I collaborated with a team of four engineers in building a referral management and sales tracking system for Dukka sales analysts. t's live at https://dexter.dukka.com. I led the Frontend React development.
The project called Dukka Referral System aggregates data from different databases in the company and computes, among other things:
- customers acquired by the company, 
- locations of customers, 
- sales, orders, and expenses of each customer, 
- top performing referees, and more. 

I found the project particularly exciting because of the data visualization aspects of the project. While I am familiar with writing scripts for cleaning data in python (using pandas), I had to write code in JavaScript to clean the data being sent from the server and put the data in formats supported by libraries like D3 and React Chart.  Additionally, in instances where I could not access Nigeria via D3, I wrote the script rendering the visualizations without leveraging any library.
The visual composition of the project was built using my skillsets across data science, web engineering, and business management. 
I have provided a link below showing some pages: https://drive.google.com/file/d/1ZQOX7nYslAnt1T3GJ9CqnFKlr66nuh4t/view?usp=drive_web
    """


def something_you_enjoy():
    return F"""
I love stand-up comedy, especially the deeply philosophical ones!
I am aware the statement sounds antithetical because the words philosophical and humorous are hardly ever used in the same sentence.
However, I think the light-heartedness with which comedians can communicate grander truths is inspiring. 
For example, Dave Chappelle's shows are a mix of humor and an exposition of the challenges faced by Black Americans, while comedians like Rick Gervais and George Carlin openly discuss atheism and explore the rationality of conventional wisdom.
Comedy is often only the scenario in which people with fixed mindsets about sensitive subjects such as racism, atheism, gender, sexism, and liberality can explore the viewpoints of others with their guard let down.  
More on this when I resume at Paystack  😉"""


def something_learnt_recently():
    return F"""I taught myself web scraping using pluralsight.com!!
Web scraping is a method of analyzing the source code of websites to pull data from the website without leveraging APIs. Web scraping can be used to collect data for sentiment analysis on social media platforms. 
I have used web scraping and web automation recently for 3 amazing things:
1. I used it to gather all property listings on propertypro.ng and then run a regression model to predict house prices. The results were fairly accurate for predicting property prices in Lagos, which was the primary project use case. 
2. I scraped all Blockchain posts on https://www.smartcontractresearch.org/ and converted them to a Pandas data frame so I can ask complicated business questions like: 
"what post had the highest number of likes while having greater than 10 comments but less than 3 tags and was written in 2022?"
3. I know Paystack uses greenhouse.io; I have written a python script to pull all job openings at 10-day intervals and send me a WhatsApp message with the json format as shown below:"""


def someting_courageous_done():
    return F"""Last year was my year of brave choices!!! I summarily skipped a career in law for a scholarship in a new school and started a company during the scholarship.
1. I skipped the Nigerian law school for a Technology Management Scholarship at the Nigerian University of Technology and Management (NUTM) - I met Shola Akinlade here, applied to this same role, passed the assessment, did the case study, got to the final stage, and was the ghosted by Paystack. 😢
2. That aside, I started a company during the program; the opportunity cost was a year of rest without the stresses of entrepreneurship. The company set out to map property transactions in Nigeria to geospatial data, thus allowing people to check the dimensions of property on a map (like Uber). It would also let consumers make land verification and certification requests from a web app and ultimately allow virtual property transfer.
In retrospect, the project failed for several reasons but mostly the disinterest of the government in releasing data on property transactions to us without physical copies of official documents. This meant our growing company, which hadn't raised any funding, would take on more risk than our underlying assets can cover in case of liability (if a document got lost). 
Thanks for making the opportunity available I look forward to speaking with the team; hopefully, Bolaji again."""


def why_company(company):
    return F""" I trust {company} because of the company's performance and the pedigree of the team members. It makes me highly comfortable with the mission and ability to deliver at {company}. I would love to work in an engineering role where I frequently interface with smart contract research and code. I intend to build my career in Blockchain (smart-contract) research. {company} is definitely one of the best places to learn from existing thought-leaders in the field."""  


def get_cover_letter_cv(template, company, role, wantCover_letter, wantCV):
    list_of_applicable_skills = "Vue, Next, ReactJS, Solidity, Flutter"
    options = "ReactJS, Python, Docker, WebGL, ExpressJS, Typescript, styled-components, React-Query, Solana, web3.js, ethers.js"
   
    front_end_cover_letter_template = F"""
The Hiring Manager,
{company}.

Hello,

{company.upper()} {role.upper()} APPLICATION.
 
I write to express interest in the {company} {role} Role. I am a mobile and web engineer at Dukka, a Nigerian FIntech company. I also have a postgraduate certification from the Nigerian University of Technology and Management (NUTM) in Data Analytics, strategy, and other business/technology-oriented courses. I was selected as one of 60 candidates across Africa for the fully funded one-year technology program.

Education:
At NUTM, I advanced my data engineering skills, I was a founding member of the Data Science Community, and I built several projects leveraging web scrapers like Beautiful Soup, cleaning the data using pandas, and running several machine learning models on sourced data. Such projects were used for stock analysis, sourcing funding opportunities and predicting Nigerian real estate prices.  I am also a recipient of the Serenze Global Technology Scholarship, where I am learning towards an AWS Developer Certification and aquiring other skills across {list_of_applicable_skills} and more.

Experience:
I have utilized my data engineering skills in several commercial scenarios. I was CTO at my prop-tech startup, where I lead a four-person team to map Geo-Spatial data in form of coordinates to property title documents in Nigeria (https://verifypro.netlify.app/). Similarly, I led the development of a web application to track a Nigerian Fintech's (Dukka) referral, sales and marketing performance (https://bit.ly/dexter-presentation). I have also written and reviewed several papers on crypto technology stacks at the Smart Contract Research Forumn in Canada.

I would like to pursue a career at {company}, a global institution where I am able to contribute to globally relevant projects amidst other thought-leaders. I am an excellent candidate for this position as I possess the skills necessary to succeed in the role and, most importantly, a curiosity to learn more.

Yours Faithfully,
    """

    venture_cover_letter_template = F"""
The Hiring Manager,
{company}.

Hello,

{company.upper()} {role.upper()} APPLICATION.
 
I write to express interest in the {company} {role} Role. I businssess and data analyst at Dukka (a Nigerian FIntech company) where I also build web and mobile apps for business analysis. 

Education:
I aquired significant business management expereince during my postgraduate program at the Nigerian University of Technology and Management (NUTM). My core learnings focused on Data Analytics, strategy, operations and more. I was selected as one of 60 candidates across Africa for the fully funded one-year technology program.
At NUTM, I also advanced my data engineering skills and I was a founding member of the Data Science Community. I built several projects to solve real world problems by writing web scrapers and running several machine learning models on sourced data. Such projects were used for stock analysis, sourcing funding opportunities and predicting Nigerian real estate prices.  I am also a recipient of the Serenze Global Technology Scholarship, where I further advanced my technical skills in solving real world problems. Lastly, I hold a Nigerian bachelor of laws degree which has among other things polished my appreciation of complex business and regulatory issues.

Experience:
I have utilized my technical skills in several commercial scenarios. I was CTO at my prop-tech startup, where I lead a four-person team to map Geo-Spatial data in form of coordinates to property title documents in Nigeria (https://verifypro.netlify.app/). Similarly, I led the development of a web application to track a Nigerian Fintech's (Dukka) referral, sales and marketing performance (https://bit.ly/dexter-presentation). I have also written and reviewed several papers on crypto technology stacks at the Smart Contract Research Forumn in Canada. Additionally, I was a research analyst at VFD Group, an investment firm where I sourced and analyzed investment opportunities in the Nigerian market and coordinated the daily relaease of our research reports. Lastly, I have also worked with accounting firms like PwC and KPMG because of a genuine need to understand global best practices.
I would like to pursue a career at {company}, a global institution where I am able to contribute to globally relevant projects amidst other thought-leaders. I am an excellent candidate for this position as I possess the skills necessary to succeed in the role and, most importantly, a curiosity to learn more.

Yours Faithfully,
    """

    business_cover_letter_template = F"""
The Hiring Manager,
{company}.

Hello,

{company.upper()} {role.upper()} APPLICATION.
 
I write to express interest in the {company} {role} Role. I businssess and data analyst at Dukka (a Nigerian FIntech company) where I also build web and mobile apps for business analysis. 

Education:
I aquired significant business management expereince during my postgraduate program at the Nigerian University of Technology and Management (NUTM). My core learnings focused on Data Analytics, strategy, operations and more. I was selected as one of 60 candidates across Africa for the fully funded one-year technology program.
At NUTM, I also advanced my data engineering skills and I was a founding member of the Data Science Community. I built several projects to solve real world problems by writing web scrapers and running several machine learning models on sourced data. Such projects were used for stock analysis, sourcing funding opportunities and predicting Nigerian real estate prices.  I am also a recipient of the Serenze Global Technology Scholarship, where I further advanced my technical skills in solving real world problems. Lastly, I hold a Nigerian bachelor of laws degree which has among other things polished my appreciation of complex business and regulatory issues.

Experience:
I have utilized my technical skills in several commercial scenarios. I was CTO at my prop-tech startup, where I lead a four-person team to map Geo-Spatial data in form of coordinates to property title documents in Nigeria (https://verifypro.netlify.app/). Similarly, I led the development of a web application to track a Nigerian Fintech's (Dukka) referral, sales and marketing performance (https://bit.ly/dexter-presentation). 
I have also written and reviewed several papers on crypto technology stacks at the Smart Contract Research Forumn. Additionally, I was a research analyst at VFD Group, an investment firm where I sourced and analyzed investment opportunities in the Nigerian market and coordinated the daily relaease of our research reports. Lastly, I have also worked with accounting firms like PwC and KPMG because of a genuine need to understand global best practices.
I would like to pursue a career at {company}, a global institution where I am able to contribute to globally relevant projects amidst other thought-leaders. I am an excellent candidate for this position as I possess the skills necessary to succeed in the role and, most importantly, a curiosity to learn more.

DeFi writing link: https://www.smartcontractresearch.org/t/research-summary-a-short-survey-on-business-models-of-decentralized-finance-defi-protocols/1947


Yours Faithfully,"""

    product_cover_letter_template = F"""        
The Hiring Manager,
{company}.

Hello,

{company.upper()} {role.upper()} APPLICATION.
 
I write to express interest in the {company} {role} Role. I businssess and data analyst at Dukka (a Nigerian FIntech company) where I also build web and mobile apps for business analysis. 

Education:
I aquired significant business management expereince during my postgraduate program at the Nigerian University of Technology and Management (NUTM). My core learnings focused on Data Analytics, strategy, operations and more. I was selected as one of 60 candidates across Africa for the fully funded one-year technology program.
At NUTM, I also advanced my data engineering skills and I was a founding member of the Data Science Community. I built several projects to solve real world problems by writing web scrapers and running several machine learning models on sourced data. Such projects were used for stock analysis, sourcing funding opportunities and predicting Nigerian real estate prices.  I am also a recipient of the Serenze Global Technology Scholarship, where I further advanced my technical skills in solving real world problems. Lastly, I hold a Nigerian bachelor of laws degree which has among other things polished my appreciation of complex business and regulatory issues.

Experience:
I have utilized my technical skills in several commercial scenarios. I was CTO at my prop-tech startup, where I lead a four-person team to map Geo-Spatial data in form of coordinates to property title documents in Nigeria (https://verifypro.netlify.app/). Similarly, I led the development of a web application to track a Nigerian Fintech's (Dukka) referral, sales and marketing performance (https://bit.ly/dexter-presentation). I have also written and reviewed several papers on crypto technology stacks at the Smart Contract Research Forumn in Canada. Additionally, I was a research analyst at VFD Group, an investment firm where I sourced and analyzed investment opportunities in the Nigerian market and coordinated the daily relaease of our research reports. Lastly, I have also worked with accounting firms like PwC and KPMG because of a genuine need to understand global best practices.
I would like to pursue a career at {company}, a global institution where I am able to contribute to globally relevant projects amidst other thought-leaders. I am an excellent candidate for this position as I possess the skills necessary to succeed in the role and, most importantly, a curiosity to learn more.

Yours Faithfully, 
    """

    crypto_hello = F"""Hello Jai, 
I am Tomide. A Nigerian software engineer with a tech and management post graduate cert. I am looking to work on globally relevant crypto-projects.
I read about Rari Capital while reading some research on yield aggregators. 
I looked you up on LinkedIn and loved your career. I am a big fan of your work on Rari Capital.
I'd love to connect, talk and work with you."""
    

    why_join = F"""I am interested in joining the DeFi/technology team. I have spent a significant amount of time writing about and reviewing DeFi papers. On the other hand, I have also spent a significant amount of time programming for web/mobile development and data analysis. Working with {company}'s technology team puts all my interests and skills in sync. """

    sendables = [front_end_cover_letter_template, "full_stack", "data_science", "data_engineer", "blockchain", business_cover_letter_template, product_cover_letter_template, venture_cover_letter_template, "mobile"]


    application_type_templates = ["front_end", "full_stack", "data_science", "data_engineer", "blockchain", "business", "product", "venture", "mobile", "technical writing"]



    def create_cover_letter():
        document = Document()
        document.add_heading('Tomide Adeoye Cover Letter', 0)
        indexPos = application_type_templates.index(template)
        document.add_paragraph(sendables[indexPos])
        document.add_paragraph(tomide_adeoye_contact)
        document_name = f"""TOMIDE ADEOYE; {company.upper()}; {role.upper()} APPLICATION COVER LETTER.docx"""
        document.save( document_name )
        return document_name

    returnObject = {
        "files": [],
        "cover_letter_text": sendables[application_type_templates.index(template)],
        "why_company": why_company,
        "topic_of_interest": topic_of_interest,
        # "cool_project": cool_project,
        "something_learnt_recently": something_learnt_recently,
    }

    if (wantCover_letter == True):
        returnObject["files"].append("/Users/tomideisawesome/Documents/GitHub/TomsWebScrapingCargo/application_maker/" + create_cover_letter())
    if (wantCV == True):
        returnObject["files"].append("/Users/tomideisawesome/Library/CloudStorage/GoogleDrive-tommideadeoye@gmail.com/My Drive/APPLICATION BOX/Tomide Adeoye CV.pdf")
    return returnObject


name, company, advice = "Denise", "BH", "Can I reach you for advice & mentorship?"
content = F"""Hello {name},
I am Tomide; a software engineer with a tech & management postgraduate certification. I am also a writer at a Technology Research Organisation. I am looking to work with amazing people on all kinds of projects.\n"""


extra = F"""Do you have any opportunities at {company} for someone like me? If not, can I still reach out for advice & mentorship and also sparingly send you projects I am building? 

Also, here is my portfolio/resume (https://tomide-adeoye.netlify.app) - it has all information about me."""


def vc_fund_template(name, company):
    return F"""
Hello {name},
My name is Tomide and I am absolutely thrilled to have been following {company}'s incredible achievements and the successes of your portfolio companies. I am eager to bring my expertise to the table and collaborate with you on devising strategies for reducing costs and increasing profitability in 2023, particularly through expanding into the African markets.

With my background in research, startups, finance, and consulting, as well as my recent postgraduate education at one of Africa's premier management and tech institutions, NUTM (nutmng.org), where I had the privilege of learning from industry leaders such as Soumitra Dutta (Dean, Saïd Business School) and Maurizio Caio (partner, TLcom Capital), I am confident in my ability to provide valuable insights and strategies. My experiences leading teams and working with organizations such as the Smart Contract Research Forum, Web3bridge, PwC, VFD Group, and startups I have consulted with have given me a unique perspective on all aspects of company operations and a deep understanding of the African Web3space. I am excited to use this knowledge to assist {company} in deal flow, portfolio management, and operations.

I have prepared a draft presentation at bit.ly/tomide-bio that provides more information about me and demonstrates how I have helped other local businesses. My contributions to {company} could include:

Utilizing analytics, data visualization, and automation to reduce costs and increase profitability,
Providing expert research and strategy for investments.

I would be honored to have the opportunity to work with {company} on a probationary basis and apply my diverse experiences. Can we schedule a time to discuss how we can achieve {company}'s 2023 objectives together?

Sincerely,
Tomide.
"""


tomide_adeoye_contact = """ \n
Tomide Adeoye.
M: +234 818 192 7251

GitHub: https://github.com/tomideadeoye
Website: https://tomide-adeoye.netlify.app
Medium: https://medium.com/@tomideadeoye/lists
LinkedIn: https://www.linkedin.com/in/tomide-adeoye-828604129
CV: https://drive.google.com/file/d/11BIzMXoQcLlLRMbb9Oa091pCSjgqOdMG/view?usp=sharing
"""

verbose_tomide_adeoye_contact = """ \n
Tomide Adeoye.
M: +234 818 192 7251

Website: https://tomide-adeoye.netlify.app
Twitter: https://twitter.com/_Tomide
GitHub: https://github.com/tomideadeoye
Medium: https://medium.com/@tomideadeoye/lists
LinkedIn: https://www.linkedin.com/in/tomide-adeoye-828604129
YouTube: https://www.youtube.com/channel/UCVQDiC3kbc6lT1JNrQGicAg

SOME PROJECT LINKS 🧑🏿‍💻:
Juice Spend: https://tomidejuiceui.netlify.app/ (React JS)
Verify Pro: https://verifypro.netlify.app/ (React JS, Firebase)
Unilever Compliance: https://unicomreport.netlify.app/ (React JS)
Referral Management: https://dexter.dukka.com/ (React JS, Django)
Book keeping mobile app: https://play.google.com/store/apps/details?id=com.dukka.dukka (Flutter, Django)
"""